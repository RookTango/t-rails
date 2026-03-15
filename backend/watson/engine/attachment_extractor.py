"""
attachment_extractor.py
Extract and compact text from change attachments before sending to Watson.

Supported formats:
  .xlsx / .xls  — Excel (primary runbook format)
  .pdf          — PDF documents
  .docx / .doc  — Word documents
  .txt .md .log .sh .py .yaml .yml .json — Plain text

Design principles:
  1. Compaction over completeness — Watson doesn't need raw cell data,
     it needs structured summaries. An Excel runbook with 500 rows becomes
     a ~800 char summary Watson can reason over.
  2. Token budget — each attachment is capped. Total across all attachments
     is also capped so a change with 5 attachments doesn't blow the prompt.
  3. Graceful degradation — if extraction fails, log it and continue.
     Never raise to the caller.
"""

import os
import logging

logger = logging.getLogger(__name__)

# Per-attachment char limit — roughly 750 tokens each
MAX_CHARS_PER_ATTACHMENT = 3000

# Total char limit across all attachments in one prompt — roughly 2000 tokens
MAX_CHARS_TOTAL = 8000


def extract_text(file_path: str, filename: str) -> str:
    """
    Extract readable text from a file.
    Returns empty string if not extractable or if extraction fails.
    """
    if not file_path or not os.path.exists(file_path):
        return ''

    ext = os.path.splitext(filename)[1].lower()

    try:
        if ext in ('.xlsx', '.xls'):
            return _extract_excel(file_path, filename)
        elif ext == '.pdf':
            return _extract_pdf(file_path)
        elif ext in ('.docx', '.doc'):
            return _extract_docx(file_path)
        elif ext in ('.txt', '.md', '.log', '.sh', '.py',
                     '.yaml', '.yml', '.json', '.csv'):
            return _extract_plain(file_path)
        else:
            return ''
    except Exception as e:
        logger.warning(f"Could not extract text from {filename}: {e}")
        return ''


# ── Excel extraction ───────────────────────────────────────────────────────

def _extract_excel(path: str, filename: str) -> str:
    """
    Extract Excel content sheet by sheet.
    Strategy:
      - Skip empty sheets and sheets with no header row
      - Convert each sheet to a compact Markdown table (max 30 rows per sheet)
      - Summarise what was found so Watson has a structural overview
      - Total output capped at MAX_CHARS_PER_ATTACHMENT
    """
    try:
        import openpyxl
    except ImportError:
        logger.warning("openpyxl not installed — pip install openpyxl")
        try:
            import xlrd
            return _extract_excel_xlrd(path)
        except ImportError:
            logger.warning("xlrd not installed either — pip install xlrd")
            return ''

    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception as e:
        logger.warning(f"Could not open Excel file {filename}: {e}")
        return ''

    sections = []
    total_chars = 0

    for sheet_name in wb.sheetnames:
        if total_chars >= MAX_CHARS_PER_ATTACHMENT:
            break

        try:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
        except Exception:
            continue

        # Skip completely empty sheets
        non_empty_rows = [r for r in rows if any(c is not None for c in r)]
        if not non_empty_rows:
            continue

        section = _sheet_to_markdown(sheet_name, non_empty_rows)
        if section:
            remaining = MAX_CHARS_PER_ATTACHMENT - total_chars
            sections.append(section[:remaining])
            total_chars += len(section)

    wb.close()

    if not sections:
        return '(Excel file contained no readable data)'

    return '\n\n'.join(sections)


def _sheet_to_markdown(sheet_name: str, rows: list) -> str:
    """
    Convert a sheet's rows to a compact Markdown table.
    - First non-empty row treated as header
    - Max 30 data rows (covers most runbooks)
    - Cells truncated to 80 chars to prevent one wide cell dominating
    - Empty columns stripped
    """
    if not rows:
        return ''

    # Find header row — first row with at least 2 non-empty cells
    header_row = None
    data_start = 0
    for i, row in enumerate(rows[:5]):
        non_empty = [c for c in row if c is not None and str(c).strip()]
        if len(non_empty) >= 2:
            header_row = row
            data_start = i + 1
            break

    if header_row is None:
        header_row = rows[0]
        data_start = 1

    # Find which columns have data
    active_cols = []
    for col_idx in range(len(header_row)):
        col_has_data = any(
            row[col_idx] is not None and str(row[col_idx]).strip() != ''
            for row in rows[data_start:data_start + 30]
            if col_idx < len(row)
        )
        if col_has_data or (header_row[col_idx] is not None):
            active_cols.append(col_idx)

    if not active_cols:
        return ''

    def cell(row, idx):
        if idx >= len(row) or row[idx] is None:
            return ''
        val = str(row[idx]).strip().replace('\n', ' ').replace('|', '/')
        return val[:80] if len(val) > 80 else val

    # Build markdown table
    lines = [f"**Sheet: {sheet_name}**"]
    headers = [cell(header_row, i) or f"Col{i+1}" for i in active_cols]
    lines.append('| ' + ' | '.join(headers) + ' |')
    lines.append('| ' + ' | '.join(['---'] * len(active_cols)) + ' |')

    data_rows = rows[data_start:data_start + 30]
    for row in data_rows:
        row_vals = [cell(row, i) for i in active_cols]
        if any(v for v in row_vals):  # skip all-empty rows
            lines.append('| ' + ' | '.join(row_vals) + ' |')

    total_rows = len(rows) - data_start
    if total_rows > 30:
        lines.append(f"*(showing 30 of {total_rows} rows)*")

    return '\n'.join(lines)


def _extract_excel_xlrd(path: str) -> str:
    """Fallback Excel extraction using xlrd (for .xls files)."""
    import xlrd
    wb = xlrd.open_workbook(path)
    sections = []
    for sheet in wb.sheets():
        rows = [sheet.row_values(i) for i in range(sheet.nrows)]
        non_empty = [r for r in rows if any(str(c).strip() for c in r)]
        if non_empty:
            section = _sheet_to_markdown(sheet.name, non_empty)
            if section:
                sections.append(section)
    return '\n\n'.join(sections)[:MAX_CHARS_PER_ATTACHMENT]


# ── PDF extraction ─────────────────────────────────────────────────────────

def _extract_pdf(path: str) -> str:
    try:
        import pdfminer.high_level
        text = pdfminer.high_level.extract_text(path)
        return _compact_text((text or '').strip())
    except ImportError:
        logger.warning("pdfminer not installed — pip install pdfminer.six")
        return ''


# ── Word document extraction ───────────────────────────────────────────────

def _extract_docx(path: str) -> str:
    try:
        import docx
        doc = docx.Document(path)
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(' | '.join(cells))
        return _compact_text('\n'.join(lines))
    except ImportError:
        logger.warning("python-docx not installed — pip install python-docx")
        return ''


# ── Plain text extraction ──────────────────────────────────────────────────

def _extract_plain(path: str) -> str:
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        return _compact_text(f.read())


# ── Text compaction ────────────────────────────────────────────────────────

def _compact_text(text: str) -> str:
    """
    Compact free-form text to stay within token budget.
    - Collapse multiple blank lines to single blank line
    - Strip lines that are just dashes/equals (decorative separators)
    - Truncate at MAX_CHARS_PER_ATTACHMENT with a marker
    """
    if not text:
        return ''

    import re
    # Collapse decorative separator lines
    text = re.sub(r'^[-=_]{3,}\s*$', '', text, flags=re.MULTILINE)
    # Collapse multiple blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.strip()

    if len(text) > MAX_CHARS_PER_ATTACHMENT:
        return text[:MAX_CHARS_PER_ATTACHMENT] + '\n... [TRUNCATED]'
    return text


# ── Main builder ───────────────────────────────────────────────────────────

def build_attachments_section(attachments: list) -> str:
    """
    Build the attachments section of the Watson prompt.

    attachments: list of dicts with keys:
      filename        — original filename
      file_path       — absolute path on disk
      attachment_type — PROCEDURE | SCREENSHOT | EVIDENCE | OTHER

    Returns a compact string ready to embed in the prompt.
    Total output is capped at MAX_CHARS_TOTAL across all attachments.
    """
    if not attachments:
        return ''

    lines = ['=== ATTACHED DOCUMENTS ===']
    total_chars = 0

    for att in attachments:
        filename  = att.get('filename', 'unknown')
        file_path = att.get('file_path', '')
        att_type  = att.get('attachment_type', '')

        if total_chars >= MAX_CHARS_TOTAL:
            lines.append(
                f"\n[{filename}] — skipped: total attachment budget exhausted"
            )
            continue

        extracted = extract_text(file_path, filename)
        budget    = MAX_CHARS_TOTAL - total_chars

        lines.append(f"\n--- {filename} ({att_type}) ---")

        if extracted:
            content = extracted[:budget]
            lines.append(content)
            total_chars += len(content)
            if len(extracted) > budget:
                lines.append('[TRUNCATED — budget reached]')
        else:
            ext = os.path.splitext(filename)[1].lower()
            if ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'):
                lines.append(
                    '(Image file — content not extractable. '
                    'Implementer should describe key findings in work notes.)'
                )
            else:
                lines.append(
                    f'(Could not extract content from {ext} file. '
                    f'Ensure required libraries are installed.)'
                )

    if len(lines) == 1:
        return ''  # Only the header — no useful content, skip entirely

    return '\n'.join(lines)